import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch

docx_path = r'c:\home\practice_2d\luxu-real-estate\luxu-real-estate\Documentacion\Documentacion_Proyecto_Base_de_Datos.docx'
pdf_path = r'c:\home\practice_2d\luxu-real-estate\luxu-real-estate\Documentacion\Documentacion_Proyecto_Base_de_Datos.pdf'
scratch = r'c:\home\practice_2d\luxu-real-estate\luxu-real-estate\scratch'

sections = [
    ('ev_panel_control.png', 'Figura 1. Módulo Panel de Control (/admin/properties)'),
    ('ev_vender.png', 'Figura 2. Formulario Vender / Agregar Inmueble (/admin/properties/add)'),
    ('ev_usuarios.png', 'Figura 3. Directorio de Usuarios y Roles (/admin/users)'),
    ('ev_comprar.png', 'Figura 4. Catálogo de Comprar (/properties?type=buy)'),
    ('ev_rentar.png', 'Figura 5. Catálogo de Rentar (/properties?type=rent)'),
    ('ev_favoritas.png', 'Figura 6. Módulo de Propiedades Favoritas (/favorites)'),
    ('ev_perfil.png', 'Figura 7. Perfil de Usuario e Historial (/profile)'),
    ('ev_agendar_visita.png', 'Figura 8. Módulo de Agendar Visita Presencial (/schedule-visit)'),
    ('ev_base_datos.png', 'Figura 9. Diagnóstico de Base de Datos (PostgreSQL 3FN + MongoDB 10,000 tuplas)')
]

# 1. BUILD DOCX
doc = docx.Document()
for s in doc.sections:
    s.top_margin = Inches(0.8)
    s.bottom_margin = Inches(0.8)
    s.left_margin = Inches(0.8)
    s.right_margin = Inches(0.8)

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_title.add_run('UNIVERSIDAD AUTÓNOMA DE YUCATÁN\nFACULTAD DE MATEMÁTICAS / INGENIERÍA DE SOFTWARE\n\n')
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0, 102, 85)

r2 = p_title.add_run('DOCUMENTACIÓN COMPLETA DE MÓDULOS Y EVIDENCIAS DEL SISTEMA\n')
r2.bold = True
r2.font.size = Pt(16)

r3 = p_title.add_run('Plataforma Luxu Real Estate - PostgreSQL 3FN & MongoDB NoSQL (10,000 Tuplas)\n\n')
r3.font.size = Pt(11)

p_intro = doc.add_paragraph('A continuación se presentan las capturas de pantalla con sesión activa de Administrador correspondientes a todos los módulos funcionales y tablas del sistema solicitadas en la rúbrica académica:')
p_intro.runs[0].font.size = Pt(10.5)

for img_name, label in sections:
    img_full = os.path.join(scratch, img_name)
    if os.path.exists(img_full):
        p_label = doc.add_paragraph()
        p_label.paragraph_format.space_before = Pt(12)
        p_label.paragraph_format.space_after = Pt(4)
        r_lbl = p_label.add_run(label)
        r_lbl.bold = True
        r_lbl.font.size = Pt(11)
        r_lbl.font.color.rgb = RGBColor(0, 102, 85)

        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_after = Pt(12)
        p_img.add_run().add_picture(img_full, width=Inches(6.2))

doc.save(docx_path)
print('DOCX saved with all 9 active session screenshots!')

# 2. BUILD PDF
doc_pdf = SimpleDocTemplate(pdf_path, pagesize=letter, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
styles = getSampleStyleSheet()
normal = styles['Normal']

title_style = ParagraphStyle(
    'PDFTitle',
    parent=normal,
    fontName='Helvetica-Bold',
    fontSize=13,
    leading=16,
    textColor=colors.HexColor('#006655'),
    alignment=1,
    spaceAfter=10
)

label_style = ParagraphStyle(
    'PDFLabel',
    parent=normal,
    fontName='Helvetica-Bold',
    fontSize=10,
    leading=13,
    textColor=colors.HexColor('#006655'),
    spaceBefore=8,
    spaceAfter=4
)

story = [
    Paragraph("UNIVERSIDAD AUTÓNOMA DE YUCATÁN - INGENIERÍA DE SOFTWARE", title_style),
    Paragraph("EVIDENCIAS DE MÓDULOS FUNCIONALES Y TABLAS DE BASE DE DATOS", title_style),
    Spacer(1, 10)
]

for img_name, label in sections:
    img_full = os.path.join(scratch, img_name)
    if os.path.exists(img_full):
        story.append(Paragraph(label, label_style))
        story.append(RLImage(img_full, width=6.5*inch, height=3.8*inch))
        story.append(Spacer(1, 10))

doc_pdf.build(story)
print('PDF saved with all 9 active session screenshots!')
