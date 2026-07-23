import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether, HRFlowable
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def generate_word_docx():
    docx_path = r"c:\home\practice_2d\luxu-real-estate\luxu-real-estate\Documentacion_Proyecto_Base_de_Datos.docx"
    doc = docx.Document()

    # Set page margins (1 inch)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x19, 0x32, 0x2F)

    # 1. PORTADA
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_univ = p_title.add_run("UNIVERSIDAD AUTÓNOMA DE YUCATÁN\nFACULTAD DE MATEMÁTICAS / INGENIERÍA DE SOFTWARE\n\n")
    run_univ.bold = True
    run_univ.font.size = Pt(14)
    run_univ.font.color.rgb = RGBColor(0x00, 0x66, 0x55)

    run_main = p_title.add_run("DOCUMENTACIÓN FINAL DE PROYECTO DE BASE DE DATOS\n")
    run_main.bold = True
    run_main.font.size = Pt(20)
    run_main.font.color.rgb = RGBColor(0x19, 0x32, 0x2F)

    run_sub = p_title.add_run("Plataforma Inmobiliaria de Lujo Luxu Real Estate\n(PostgreSQL + MongoDB + Redis + Next.js 16)\n\n\n")
    run_sub.font.size = Pt(14)
    run_sub.font.color.rgb = RGBColor(0x5C, 0x70, 0x6D)

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.line_spacing = 1.3
    
    meta_text = """
    ASIGNATURA: Proyecto de Base de Datos / Gestión de Bases de Datos
    ESTUDIANTE: Santiago Asahel Pech
    MATRÍCULA: 2026-BD-9941
    PROFESOR / EVALUADOR: Comité Académico de Bases de Datos
    SEMESTRE: 2026-I
    FECHA DE ENTREGA: 22 de Julio de 2026
    PUNTAJE OBJETIVO: 100 / 100 Puntos
    """
    run_meta = p_meta.add_run(meta_text)
    run_meta.font.size = Pt(11)
    run_meta.bold = True
    run_meta.font.color.rgb = RGBColor(0x19, 0x32, 0x2F)

    doc.add_page_break()

    # Helper function for headings
    def add_custom_heading(text, level):
        h = doc.add_heading(text, level=level)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)
        run = h.runs[0]
        if level == 1:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0x00, 0x66, 0x55)
            run.bold = True
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x19, 0x32, 0x2F)
            run.bold = True
        elif level == 3:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x5C, 0x70, 0x6D)
            run.bold = True
        return h

    # 2. DESARROLLO
    add_custom_heading("DESARROLLO (96 PUNTOS)", level=1)

    # 2.1 Descripción del problema
    add_custom_heading("1. Descripción del Problema (3 pts.)", level=2)
    p = doc.add_paragraph()
    p.add_run("""El sector inmobiliario de alta gama presenta requerimientos estrictos de gestión de datos, confidencialidad y tiempos de respuesta inmediatos. En los entornos tradicionales de comercialización inmobiliaria se identifican los siguientes problemas estructurales:

1. Desorganización y Redundancia de Información: Las fichas técnicas de las propiedades, los datos de contacto de los clientes VIP y las agendas de los agentes se gestionan comúnmente en hojas de cálculo independientes. Esto ocasiona inconsistencia en los datos, duplicidad de inmuebles y pérdida de oportunidades de venta.
2. Ineficiencia en Consultas Complejas: Los compradores exigentes requieren filtrar inmuebles bajo combinaciones estrictas de variables (rango de precio, ubicación exacta, número de habitaciones, baños, tipo de inmueble y amenidades exclusivas como heli-puerto, alberca privada o cava de vinos). Sin una base de datos relacional adecuadamente estructurada en 3FN e indexada, las consultas tardan segundos excesivos en responder.
3. Desafíos de Escalabilidad y Analítica Masiva: La necesidad de registrar miles de interacciones diarias (vistas de páginas, clics en favoritos y registros de eventos) satura los motores transaccionales si no se cuenta con un esquema híbrido NoSQL diseñado para analítica masiva.
4. Falta de Control de Acceso basado en Roles (RBAC): Los sistemas vulnerables carecen de segregación entre usuarios administradores y clientes, lo que permite alteraciones no autorizadas en precios o estados comerciales.

Solución Implementada: La plataforma Luxu Real Estate resuelve de forma integral esta problemática mediante una arquitectura híbrida de alto rendimiento que combina PostgreSQL (BD Relacional para transacciones ACID), MongoDB (BD NoSQL para analítica documental masiva) y Redis (BD Clave-Valor para almacenamiento en caché de alto rendimiento).""")

    # 2.2 Alcance del proyecto
    add_custom_heading("2. Alcance del Proyecto (4 pts.)", level=2)
    p = doc.add_paragraph()
    p.add_run("""El proyecto abarca el desarrollo completo de la capa de almacenamiento y la capa de aplicación web responsiva.

A. Catálogos con Operaciones CRUD y Búsquedas:
- Catálogo de Propiedades: Operaciones completas de Agregar (Create), Consultar y Filtrar (Read), Editar datos técnicos/precios (Update) y Borrar inmuebles (Delete). Incluye búsqueda inteligente por términos clave.
- Catálogo de Usuarios y Perfiles: Gestión de usuarios con control de roles (Admin, Agente, Cliente) y datos de perfil.
- Catálogo de Favoritos: Gestión de inmuebles guardados con persistencia en localStorage y sincronización en la base de datos relacional.
- Catálogo de Citas y Visitas Guiadas: Agendado de recorridos presenciales conectando clientes con propiedades.

B. Módulos Especializados:
- Módulo de Control de Acceso por Roles (RBAC): Filtra automáticamente la interfaz de usuario en la barra de navegación según el rol autenticado (🛡️ Admin vs 👤 Cliente).
- Módulo de KPIs y Analítica Comercial: Panel de métricas en tiempo real que calcula propiedades activas, en renta, en venta y pendientes de aprobación.""")

    # 2.3 Modelo Relacional y Normalización
    add_custom_heading("3. Modelo de Base de Datos Relacional y Normalización (1FN, 2FN, 3FN) (5 pts.)", level=2)
    p = doc.add_paragraph()
    p.add_run("""La normalización es el proceso formal para eliminar redundancias y anomalías de inserción, actualización y borrado en bases de datos relacionales.

Fundamentos Teóricos:
- Primera Forma Normal (1FN): Requiere que los atributos sean atómicos (sin valores compuestos ni repetidos) y que exista una clave primaria única.
- Segunda Forma Normal (2FN): Exige estar en 1FN y que todos los atributos no clave dependan funcionalmente de la clave primaria completa (sin dependencias parciales).
- Tercera Forma Normal (3FN): Exige estar en 2FN y que ningún atributo no clave dependa de otro atributo no clave (sin dependencias transitivas).

Proceso de Aplicación Paso a Paso en Luxu Real Estate:""")

    # Table for Normalization
    table_norm = doc.add_table(rows=5, cols=3)
    table_norm.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table_norm.rows[0].cells
    hdr[0].text = "Fase de Normalización"
    hdr[1].text = "Estructura de la Tabla"
    hdr[2].text = "Acción / Justificación Aplicada"
    for cell in hdr:
        set_cell_background(cell, "006655")
        for p in cell.paragraphs:
            p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            p.runs[0].font.bold = True

    rows_norm_data = [
        ("No Normalizada (UNF)", "TABLA_REGISTRO(ID_Prop, Titulo, Precio, Dueno_Nombre, Dueno_Email, Amenidades_Lista)", "Contiene datos repetitivos y amenidades en lista atómica en una sola celda."),
        ("1ª Forma Normal (1FN)", "PROPIEDAD(ID_Prop, Titulo, Precio, Dueno_Nombre, Dueno_Email), AMENIDAD(ID_Prop, Amenidad)", "Se eliminan grupos repetitivos garantizando atomicidad de valores."),
        ("2ª Forma Normal (2FN)", "PROPIEDAD(ID_Prop, Titulo, Precio, Owner_ID), PROPIETARIO(Owner_ID, Nombre, Email)", "Se eliminan dependencias parciales separando los datos del propietario."),
        ("3ª Forma Normal (3FN)", "profiles(id, full_name, email, role)\nproperties(id, slug, title, price, owner_id)\nuser_favorites(id, user_id, property_id)\nappointments(id, user_id, property_id)", "Se eliminan dependencias transitivas. Todos los atributos dependen únicamente de la PK.")
    ]

    for idx, (fase, est, acc) in enumerate(rows_norm_data):
        row_cells = table_norm.rows[idx+1].cells
        row_cells[0].text = fase
        row_cells[1].text = est
        row_cells[2].text = acc
        if idx % 2 == 1:
            for cell in row_cells:
                set_cell_background(cell, "F5F7F7")

    doc.add_paragraph()

    # 2.4 Modelo No Relacional
    add_custom_heading("4. Modelo de Base de Datos No Relacional (10 pts.)", level=2)
    p = doc.add_paragraph()
    p.add_run("""Se diseñó un modelo no relacional utilizando MongoDB (Orientado a Documentos) y Redis (Clave-Valor):

1. Modelo Orientado a Documentos (MongoDB Atlas):
Almacena fichas enriquecidas de propiedades con esquemas dinámicos, incluyendo coordenadas geográficas, colecciones de imágenes, amenidades y estadísticas de visualización integradas en un solo documento BSON sin requerir operaciones JOIN costosas.

Estructura del Documento BSON en MongoDB:""")

    p_code = doc.add_paragraph()
    p_code.paragraph_format.left_indent = Inches(0.5)
    r_code = p_code.add_run("""{
  "_id": "66a01b2c00000001",
  "slug": "propiedad-casa-beverly-hills-1",
  "title": "Casa Lujosa #1 en Beverly Hills",
  "price": 5250000,
  "status": "ACTIVE",
  "specs": { "beds": 5, "baths": 4.5, "sqft": 4200, "garage_spaces": 3 },
  "location": {
    "city_state": "Beverly Hills, CA",
    "address": "742 Luxury Avenue",
    "coordinates": { "lat": 34.0736, "lng": -118.4004 }
  },
  "features": {
    "amenities": ["Alberca Privada", "Casa Inteligente", "Vista al Mar", "Gimnasio Privado"],
    "year_built": 2024
  },
  "analytics": { "views": 1420, "saved_in_favorites_count": 89 }
}""")
    r_code.font.name = 'Consolas'
    r_code.font.size = Pt(9.5)
    r_code.font.color.rgb = RGBColor(0x19, 0x32, 0x2F)

    # 2.5 Manejo de Restricciones
    add_custom_heading("5. Manejo de Restricciones (5 pts.)", level=2)
    p = doc.add_paragraph()
    p.add_run("El esquema PostgreSQL aplica estrictamente las 6 restricciones de integridad relacional:")

    table_rest = doc.add_table(rows=7, cols=4)
    table_rest.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_r = table_rest.rows[0].cells
    hdr_r[0].text = "Restricción"
    hdr_r[1].text = "Tabla / Campo"
    hdr_r[2].text = "Sintaxis SQL"
    hdr_r[3].text = "Propósito"
    for cell in hdr_r:
        set_cell_background(cell, "006655")
        for p in cell.paragraphs:
            p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            p.runs[0].font.bold = True

    rest_data = [
        ("PRIMARY KEY", "profiles.id, properties.id", "id UUID PRIMARY KEY DEFAULT gen_random_uuid()", "Garantiza unicidad de clave primaria."),
        ("FOREIGN KEY", "properties.owner_id", "REFERENCES profiles(id) ON DELETE SET NULL", "Asegura integridad referencial."),
        ("UNIQUE", "profiles.email, properties.slug", "email VARCHAR(150) UNIQUE NOT NULL", "Impide correos o slugs duplicados."),
        ("NOT NULL", "properties.title, properties.price", "title VARCHAR(200) NOT NULL", "Evita datos faltantes o nulos."),
        ("CHECK", "properties.price, profiles.role", "CHECK (price >= 0), CHECK (role IN (...))", "Valida rangos y dominio de valores."),
        ("DEFAULT", "profiles.role, properties.status", "role DEFAULT 'Cliente'", "Asigna valor por defecto automático.")
    ]

    for idx, (res, tab, syn, prp) in enumerate(rest_data):
        row_cells = table_rest.rows[idx+1].cells
        row_cells[0].text = res
        row_cells[1].text = tab
        row_cells[2].text = syn
        row_cells[3].text = prp
        if idx % 2 == 1:
            for cell in row_cells:
                set_cell_background(cell, "F5F7F7")

    doc.add_paragraph()

    # 2.6 Diccionario de Datos
    add_custom_heading("6. Diccionario de Datos (10 pts.)", level=2)
    p = doc.add_paragraph()
    p.add_run("Estructura detallada de las tablas relacionales y colecciones NoSQL:")

    # Diccionario Tabla Properties
    add_custom_heading("Diccionario de Tabla Relacional: properties", level=3)
    t_dic = doc.add_table(rows=11, cols=7)
    t_dic.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_d = t_dic.rows[0].cells
    hdr_d[0].text = "Campo"
    hdr_d[1].text = "Tipo Dato"
    hdr_d[2].text = "Long."
    hdr_d[3].text = "Llave"
    hdr_d[4].text = "Nulo"
    hdr_d[5].text = "Descripción"
    hdr_d[6].text = "Restricciones"
    for cell in hdr_d:
        set_cell_background(cell, "006655")
        for p in cell.paragraphs:
            p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            p.runs[0].font.bold = True

    dic_props_data = [
        ("id", "UUID", "36", "PK", "NO", "Identificador único", "gen_random_uuid()"),
        ("slug", "VARCHAR", "200", "UQ", "NO", "URL amigable de la propiedad", "UNIQUE"),
        ("title", "VARCHAR", "200", "-", "NO", "Nombre de la residencia", "NOT NULL"),
        ("price", "NUMERIC", "15,2", "-", "NO", "Precio publicado en USD", "CHECK (price >= 0)"),
        ("status", "VARCHAR", "20", "-", "NO", "Estado (ACTIVE, FOR SALE, FOR RENT)", "CHECK status IN (...)"),
        ("beds", "INT", "4", "-", "NO", "Número de recámaras", "CHECK (beds >= 0)"),
        ("baths", "NUMERIC", "3,1", "-", "NO", "Número de baños", "DEFAULT 1.0"),
        ("sqft", "INT", "4", "-", "NO", "Superficie en m²", "CHECK (sqft > 0)"),
        ("location", "VARCHAR", "150", "-", "NO", "Ciudad y Estado", "NOT NULL"),
        ("owner_id", "UUID", "36", "FK", "SI", "ID del usuario propietario", "REFERENCES profiles(id)")
    ]

    for idx, (c, td, l, k, n, d, r) in enumerate(dic_props_data):
        row = t_dic.rows[idx+1].cells
        row[0].text = c
        row[1].text = td
        row[2].text = l
        row[3].text = k
        row[4].text = n
        row[5].text = d
        row[6].text = r
        if idx % 2 == 1:
            for cell in row:
                set_cell_background(cell, "F5F7F7")

    doc.add_paragraph()

    # 2.7 Creación de las BD
    add_custom_heading("7. Creación de las Bases de Datos (10 pts.)", level=2)
    p = doc.add_paragraph()
    p.add_run("El proyecto cuenta con scripts automáticos DDL y DML para PostgreSQL y scripts de inicialización en MongoDB:")

    # 2.8 Inserción de Registros
    add_custom_heading("8. Inserción de Registros (15 pts.)", level=2)
    p = doc.add_paragraph()
    p.add_run("""Cumpliendo con los requisitos de volumen de datos:
1. Base de Datos Relacional (PostgreSQL): Se generaron e insertaron 50 registros reales y coherentes por cada una de las 4 tablas (profiles, properties, user_favorites, appointments), sumando 200 tuplas relacionales.
2. Base de Datos No Relacional (MongoDB): Se generaron e insertaron 10,000 documentos BSON completos con coordenadas, analítica y especificidades técnicas para pruebas de carga masiva.""")

    # 2.9 Archivos Adjuntos y Backup
    add_custom_heading("9. Archivos Adjuntos y Copia de Seguridad (4 pts.)", level=2)
    p = doc.add_paragraph()
    p.add_run("""Se incluyen en el entregable del proyecto los siguientes archivos ejecutables de respaldo:
- schema_and_data_postgresql.sql: Script completo DDL de creación de tablas, llaves, restricciones e inserción de 50 registros por tabla.
- mongodb_nosql_dataset.json: Dataset masivo JSON de 10,000 registros para importar directamente en MongoDB Atlas con mongoimport.""")

    # 2.10 Proyecto Web
    add_custom_heading("10. Proyecto Web Desarrollado con Conexión a BD (20 pts.)", level=2)
    p = doc.add_paragraph()
    p.add_run("""La aplicación web fue construida utilizando la arquitectura moderna de desarrollo de software de Next.js 16 (App Router), React 19 y TypeScript.

Características de la Arquitectura:
- Conexión a PostgreSQL: Mediante el cliente oficial de Supabase SDK, ejecutando consultas seguras.
- Conexión NoSQL / Redis: Integración de cliente Redis para manejo de caché de propiedades en memoria.
- Interfaz Gráfica de Alto Impacto: Construida con Tailwind CSS 4, componentes modulares y soporte multilenguaje (Español / Inglés).""")

    # 2.11 Identificación de Objetos
    add_custom_heading("11. Identificación de Objetos en la Base de Datos (10 pts.)", level=2)
    p = doc.add_paragraph()
    p.add_run("Matriz de objetos programados y su ubicación en el motor de base de datos:")

    t_obj = doc.add_table(rows=7, cols=5)
    t_obj.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_o = t_obj.rows[0].cells
    hdr_o[0].text = "Objeto"
    hdr_o[1].text = "Nombre BD"
    hdr_o[2].text = "Función Operativa"
    hdr_o[3].text = "Tablas Afectadas"
    hdr_o[4].text = "Ubicación en el Motor"
    for cell in hdr_o:
        set_cell_background(cell, "006655")
        for p in cell.paragraphs:
            p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            p.runs[0].font.bold = True

    obj_data = [
        ("Tabla Base", "properties", "Guarda el catálogo maestro de inmuebles", "properties", "Esquema public / pg_class"),
        ("Vista", "vw_active_properties", "Proporciona resumen acelerado de inmuebles activos", "properties", "Catálogo pg_views"),
        ("Disparador", "trg_update_timestamp", "Actualiza automáticamente fecha de modificación", "properties, profiles", "Catálogo pg_trigger"),
        ("Función", "fn_calculate_kpis()", "Calcula métricas de propiedades activas y en renta", "properties", "Catálogo pg_proc"),
        ("Procedimiento", "sp_schedule_visit()", "Ejecuta transaccionalmente el agendado de cita", "appointments", "Catálogo pg_proc"),
        ("Índice B-Tree", "idx_prop_slug", "Acelera búsquedas por URL amigable en O(log N)", "properties", "Catálogo pg_am")
    ]

    for idx, (o, n, f, t, u) in enumerate(obj_data):
        row = t_obj.rows[idx+1].cells
        row[0].text = o
        row[1].text = n
        row[2].text = f
        row[3].text = t
        row[4].text = u
        if idx % 2 == 1:
            for cell in row:
                set_cell_background(cell, "F5F7F7")

    doc.add_paragraph()

    # 2.12 Funcionalidad de la Aplicación
    add_custom_heading("12. Funcionalidad de la Aplicación con Capturas (10 pts.)", level=2)
    p = doc.add_paragraph()
    p.add_run("""Muestra detallada del sistema en funcionamiento:

1. Pantalla Principal (Landing Page): Buscador de propiedades con selección rápida por categorías (Casas, Departamentos, Villas, Penthouses), catálogo traducido al español e integración del modal de filtros.
2. Catálogo General de Propiedades (/properties): Vista de cuadrícula responsiva con tarjetas detallando recámaras, baños, superficie en m², insignias de estado (EN RENTA / EXCLUSIVO) y botón independiente de favoritas.
3. Vista de Favoritos (/favorites): Módulo que sincroniza en tiempo real las propiedades guardadas por el usuario, con ordenamiento por precio y vista en lista o cuadrícula.
4. Panel Administrativo (/admin/properties): Dashboard de gestión con KPIs dinámicos, tabla interactiva de inmuebles y gestión de estados.""")

    # 3. REFERENCIAS
    doc.add_page_break()
    add_custom_heading("REFERENCIAS EN FORMATO APA 7ª EDICIÓN (2 PUNTOS)", level=1)
    
    p_ref = doc.add_paragraph()
    p_ref.paragraph_format.line_spacing = 1.3
    p_ref.paragraph_format.left_indent = Inches(0.5)
    p_ref.paragraph_format.first_line_indent = Inches(-0.5)

    refs = [
        "Chodorow, C. (2013). MongoDB: The Definitive Guide (2.ª ed.). O'Reilly Media.",
        "Date, C. J. (2004). An Introduction to Database Systems (8.ª ed.). Addison-Wesley.",
        "Elmasri, R., & Navathe, S. B. (2017). Fundamentos de Sistemas de Bases de Datos (7.ª ed.). Pearson Educación. (Libro de texto principal)",
        "PostgreSQL Global Development Group. (2026). PostgreSQL 16.0 Documentation. PostgreSQL.org. https://www.postgresql.org/docs/16/",
        "Silberschatz, A., Korth, H. F., & Sudarshan, S. (2020). Database System Concepts (7.ª ed.). McGraw-Hill Education."
    ]

    for ref in refs:
        doc.add_paragraph(ref)

    doc.save(docx_path)
    print(f"Word DOCX documentation created successfully at {docx_path}")

def generate_pdf_doc():
    pdf_path = r"c:\home\practice_2d\luxu-real-estate\luxu-real-estate\Documentacion_Proyecto_Base_de_Datos.pdf"
    doc = SimpleDocTemplate(
        pdf_path,
        pagesize=letter,
        rightMargin=0.75*inch, leftMargin=0.75*inch,
        topMargin=0.75*inch, bottomMargin=0.75*inch
    )

    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle(
        'CoverTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=20,
        leading=24,
        textColor=colors.HexColor('#19322F'),
        alignment=1, # Center
        spaceAfter=15
    )

    subtitle_style = ParagraphStyle(
        'CoverSubtitle',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=13,
        leading=16,
        textColor=colors.HexColor('#006655'),
        alignment=1,
        spaceAfter=30
    )

    h1_style = ParagraphStyle(
        'Heading1Custom',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=16,
        leading=20,
        textColor=colors.HexColor('#006655'),
        spaceBefore=18,
        spaceAfter=10,
        keepWithNext=True
    )

    h2_style = ParagraphStyle(
        'Heading2Custom',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=12,
        leading=15,
        textColor=colors.HexColor('#19322F'),
        spaceBefore=12,
        spaceAfter=6,
        keepWithNext=True
    )

    body_style = ParagraphStyle(
        'BodyCustom',
        parent=styles['BodyText'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        textColor=colors.HexColor('#19322F'),
        spaceAfter=8
    )

    code_style = ParagraphStyle(
        'CodeCustom',
        parent=styles['Code'],
        fontName='Courier',
        fontSize=8.5,
        leading=11,
        textColor=colors.HexColor('#19322F'),
        backColor=colors.HexColor('#F5F7F7'),
        borderColor=colors.HexColor('#006655'),
        borderWidth=0.5,
        borderPadding=6,
        spaceAfter=10
    )

    story = []

    # Cover Page
    story.append(Spacer(1, 0.5*inch))
    story.append(Paragraph("UNIVERSIDAD AUTÓNOMA DE YUCATÁN", ParagraphStyle('Univ', parent=subtitle_style, fontSize=14, leading=17, fontName='Helvetica-Bold')))
    story.append(Paragraph("FACULTAD DE MATEMÁTICAS / INGENIERÍA DE SOFTWARE", ParagraphStyle('Fac', parent=subtitle_style, fontSize=11, leading=14)))
    story.append(Spacer(1, 0.8*inch))
    story.append(Paragraph("DOCUMENTACIÓN FINAL DE PROYECTO DE BASE DE DATOS", title_style))
    story.append(Paragraph("Plataforma Inmobiliaria de Lujo Luxu Real Estate<br/>(PostgreSQL + MongoDB + Redis + Next.js 16)", subtitle_style))
    story.append(Spacer(1, 1.2*inch))

    meta_pdf = """
    <b>ASIGNATURA:</b> Proyecto de Base de Datos / Gestión de Bases de Datos<br/>
    <b>ESTUDIANTE:</b> Santiago Asahel Pech<br/>
    <b>MATRÍCULA:</b> 2026-BD-9941<br/>
    <b>PROFESOR / EVALUADOR:</b> Comité Académico de Bases de Datos<br/>
    <b>SEMESTRE:</b> 2026-I<br/>
    <b>FECHA DE ENTREGA:</b> 22 de Julio de 2026<br/>
    <b>PUNTAJE OBJETIVO:</b> 100 / 100 Puntos
    """
    story.append(Paragraph(meta_pdf, ParagraphStyle('Meta', parent=body_style, alignment=1, leading=18, fontSize=11)))
    story.append(PageBreak())

    # Section 1: Development
    story.append(Paragraph("DESARROLLO (96 PUNTOS)", h1_style))
    story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor('#006655'), spaceAfter=15))

    story.append(Paragraph("1. Descripción del Problema (3 pts.)", h2_style))
    p1 = ("El sector inmobiliario de alta gama requiere un manejo riguroso de información de alto rendimiento. "
          "La plataforma <b>Luxu Real Estate</b> resuelve los cuellos de botella de desorganización de datos, "
          "consultas lentas y falta de control de acceso mediante un esquema de base de datos híbrido relacional y NoSQL.")
    story.append(Paragraph(p1, body_style))

    story.append(Paragraph("2. Alcance del Proyecto (4 pts.)", h2_style))
    p2 = ("El proyecto abarca catálogos CRUD para Propiedades, Usuarios, Favoritos y Citas de Visitas, "
          "además de módulos especializados de Control de Acceso por Roles (RBAC) y Métricas de Negocio (KPIs).")
    story.append(Paragraph(p2, body_style))

    story.append(Paragraph("3. Modelo Relacional y Normalización (1FN, 2FN, 3FN) (5 pts.)", h2_style))
    story.append(Paragraph("Aplicación de las 3 Formas Normales para eliminar redundancias y dependencias transitivas:", body_style))

    data_t1 = [
        [Paragraph("<b>Fase Normalización</b>", body_style), Paragraph("<b>Estructura Aplicada</b>", body_style), Paragraph("<b>Resultado / Justificación</b>", body_style)],
        [Paragraph("1FN", body_style), Paragraph("Atomicidad de atributos", body_style), Paragraph("Eliminación de amenidades multivaluadas.", body_style)],
        [Paragraph("2FN", body_style), Paragraph("Eliminación dep. parciales", body_style), Paragraph("Separación de propietario a tabla profiles.", body_style)],
        [Paragraph("3FN (Final)", body_style), Paragraph("profiles, properties, user_favorites, appointments", body_style), Paragraph("Todos los atributos dependen únicamente de la PK.", body_style)]
    ]
    t1 = Table(data_t1, colWidths=[1.2*inch, 2.5*inch, 3.1*inch])
    t1.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#006655')),
        ('TEXTCOLOR', (0,0), (-1,0), colors.white),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#CCCCCC')),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#F5F7F7')])
    ]))
    story.append(t1)
    story.append(Spacer(1, 10))

    story.append(Paragraph("4. Modelo No Relacional (10 pts.)", h2_style))
    story.append(Paragraph("Ejemplo de documento BSON en MongoDB Atlas para analítica masiva:", body_style))
    json_sample = """{
  "_id": "66a01b2c00000001",
  "slug": "propiedad-casa-beverly-hills-1",
  "title": "Casa Lujosa #1 en Beverly Hills",
  "price": 5250000,
  "status": "ACTIVE",
  "specs": { "beds": 5, "baths": 4.5, "sqft": 4200 },
  "analytics": { "views": 1420, "saved_in_favorites_count": 89 }
}"""
    story.append(Paragraph(json_sample.replace('\n', '<br/>').replace(' ', '&nbsp;'), code_style))

    story.append(Paragraph("5. Manejo de Restricciones (5 pts.)", h2_style))
    data_t2 = [
        [Paragraph("<b>Restricción</b>", body_style), Paragraph("<b>Tabla/Campo</b>", body_style), Paragraph("<b>Código SQL</b>", body_style)],
        [Paragraph("PRIMARY KEY", body_style), Paragraph("profiles.id, properties.id", body_style), Paragraph("id UUID PRIMARY KEY DEFAULT gen_random_uuid()", body_style)],
        [Paragraph("FOREIGN KEY", body_style), Paragraph("properties.owner_id", body_style), Paragraph("REFERENCES profiles(id) ON DELETE SET NULL", body_style)],
        [Paragraph("UNIQUE", body_style), Paragraph("profiles.email", body_style), Paragraph("email VARCHAR(150) UNIQUE NOT NULL", body_style)],
        [Paragraph("CHECK", body_style), Paragraph("properties.price", body_style), Paragraph("CHECK (price >= 0)", body_style)]
    ]
    t2 = Table(data_t2, colWidths=[1.5*inch, 2.3*inch, 3.0*inch])
    t2.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#006655')),
        ('TEXTCOLOR', (0,0), (-1,0), colors.white),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#CCCCCC')),
        ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#F5F7F7')])
    ]))
    story.append(t2)
    story.append(Spacer(1, 10))

    story.append(Paragraph("6. Diccionario de Datos (10 pts.)", h2_style))
    story.append(Paragraph("7. Creación de las BD Relacional y NoSQL (10 pts.)", h2_style))
    story.append(Paragraph("8. Inserción de 50 Registros (Relacional) y 10,000 Registros (NoSQL) (15 pts.)", h2_style))
    story.append(Paragraph("Se generaron 200 tuplas relacionales en PostgreSQL y 10,000 documentos BSON en MongoDB.", body_style))

    story.append(Paragraph("9. Archivos Adjuntos y Backup (4 pts.)", h2_style))
    story.append(Paragraph("Archivos adjuntos: <b>schema_and_data_postgresql.sql</b> y <b>mongodb_nosql_dataset.json</b>.", body_style))

    story.append(Paragraph("10. Proyecto Web Desarrollado con Conexión a BD (20 pts.)", h2_style))
    story.append(Paragraph("Desarrollado en Next.js 16 + React 19 + TypeScript + Supabase PostgreSQL + Redis.", body_style))

    story.append(Paragraph("11. Identificación de Objetos en la Base de Datos (10 pts.)", h2_style))
    data_t3 = [
        [Paragraph("<b>Objeto</b>", body_style), Paragraph("<b>Nombre BD</b>", body_style), Paragraph("<b>Función Operativa</b>", body_style), Paragraph("<b>Ubicación Motor</b>", body_style)],
        [Paragraph("Tabla", body_style), Paragraph("properties", body_style), Paragraph("Catálogo maestro de inmuebles", body_style), Paragraph("Schema public", body_style)],
        [Paragraph("Vista", body_style), Paragraph("vw_active_properties", body_style), Paragraph("Resumen acelerado de inmuebles activos", body_style), Paragraph("pg_views", body_style)],
        [Paragraph("Trigger", body_style), Paragraph("trg_update_timestamp", body_style), Paragraph("Actualiza fecha de modificación", body_style), Paragraph("pg_trigger", body_style)],
        [Paragraph("Índice", body_style), Paragraph("idx_prop_slug", body_style), Paragraph("Búsqueda acelerada por URL en O(log N)", body_style), Paragraph("pg_am", body_style)]
    ]
    t3 = Table(data_t3, colWidths=[1.1*inch, 1.6*inch, 2.5*inch, 1.6*inch])
    t3.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#006655')),
        ('TEXTCOLOR', (0,0), (-1,0), colors.white),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#CCCCCC')),
        ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#F5F7F7')])
    ]))
    story.append(t3)
    story.append(Spacer(1, 10))

    story.append(Paragraph("12. Funcionalidad de la Aplicación (10 pts.)", h2_style))
    story.append(Paragraph("Demostración con capturas de la interfaz web responsiva, buscador en tiempo real, catálogo de favoritas y panel administrativo RBAC.", body_style))

    story.append(PageBreak())
    story.append(Paragraph("REFERENCIAS EN FORMATO APA 7ª EDICIÓN (2 PUNTOS)", h1_style))
    story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor('#006655'), spaceAfter=15))

    refs = [
        "Chodorow, C. (2013). <i>MongoDB: The Definitive Guide</i> (2.ª ed.). O'Reilly Media.",
        "Date, C. J. (2004). <i>An Introduction to Database Systems</i> (8.ª ed.). Addison-Wesley.",
        "Elmasri, R., & Navathe, S. B. (2017). <i>Fundamentos de Sistemas de Bases de Datos</i> (7.ª ed.). Pearson Educación.",
        "PostgreSQL Global Development Group. (2026). <i>PostgreSQL 16.0 Documentation</i>. PostgreSQL.org. https://www.postgresql.org/docs/16/",
        "Silberschatz, A., Korth, H. F., & Sudarshan, S. (2020). <i>Database System Concepts</i> (7.ª ed.). McGraw-Hill Education."
    ]

    for ref in refs:
        story.append(Paragraph(ref, ParagraphStyle('RefStyle', parent=body_style, leftIndent=20, firstLineIndent=-20, spaceAfter=8)))

    doc.build(story)
    print(f"PDF documentation created successfully at {pdf_path}")

generate_word_docx()
generate_pdf_doc()
