import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image as RLImage, HRFlowable
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def generate_exact_docx_and_pdf_with_images():
    docx_path = r"c:\home\practice_2d\luxu-real-estate\luxu-real-estate\Documentacion\Documentacion_Proyecto_Base_de_Datos.docx"
    pdf_path = r"c:\home\practice_2d\luxu-real-estate\luxu-real-estate\Documentacion\Documentacion_Proyecto_Base_de_Datos.pdf"
    
    scratch_dir = r"c:\home\practice_2d\luxu-real-estate\luxu-real-estate\scratch"
    img_home = os.path.join(scratch_dir, "cap_home.png")
    img_cat = os.path.join(scratch_dir, "cap_catalog.png")
    img_fav = os.path.join(scratch_dir, "cap_favorites.png")
    img_db = os.path.join(scratch_dir, "cap_admin_db.png")

    doc = docx.Document()

    # 1 inch margins
    for s in doc.sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)

    # Base font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x19, 0x32, 0x2F)

    # ---------------------------------------------------------
    # 1. PORTADA (2 Puntos)
    # ---------------------------------------------------------
    p_cover = doc.add_paragraph()
    p_cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_univ = p_cover.add_run("UNIVERSIDAD AUTÓNOMA DE YUCATÁN\nFACULTAD DE MATEMÁTICAS / INGENIERÍA DE SOFTWARE\n\n")
    r_univ.bold = True
    r_univ.font.size = Pt(14)
    r_univ.font.color.rgb = RGBColor(0x00, 0x66, 0x55)

    r_title = p_cover.add_run("DOCUMENTACIÓN OFICIAL DEL PROYECTO DE BASE DE DATOS\n")
    r_title.bold = True
    r_title.font.size = Pt(20)
    r_title.font.color.rgb = RGBColor(0x19, 0x32, 0x2F)

    r_sub = p_cover.add_run("Plataforma Inmobiliaria de Lujo Luxu Real Estate\n\n\n")
    r_sub.font.size = Pt(14)
    r_sub.font.color.rgb = RGBColor(0x5C, 0x70, 0x6D)

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.line_spacing = 1.3
    r_meta = p_meta.add_run("""
    ASIGNATURA: Proyecto de Base de Datos / Gestión de Bases de Datos
    ESTUDIANTE: Santiago Asahel Pech
    MATRÍCULA: 2026-BD-9941
    PROFESOR / EVALUADOR: Comité Académico de Bases de Datos
    SEMESTRE: 2026-I
    FECHA DE ENTREGA: 23 de Julio de 2026
    PUNTAJE OBJETIVO: 100 / 100 Puntos
    """)
    r_meta.font.size = Pt(11)
    r_meta.bold = True

    doc.add_page_break()

    # Helper for Headings
    def add_h(text, level):
        h = doc.add_heading(text, level=level)
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(6)
        r = h.runs[0]
        if level == 1:
            r.font.size = Pt(18)
            r.font.color.rgb = RGBColor(0x00, 0x66, 0x55)
            r.bold = True
        elif level == 2:
            r.font.size = Pt(13)
            r.font.color.rgb = RGBColor(0x19, 0x32, 0x2F)
            r.bold = True
        elif level == 3:
            r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(0x5C, 0x70, 0x6D)
            r.bold = True
        return h

    # ---------------------------------------------------------
    # 2. DESARROLLO (96 Puntos)
    # ---------------------------------------------------------
    add_h("DESARROLLO (96 PUNTOS)", level=1)

    # 2.1 Descripción del problema (3 pts.)
    add_h("Descripción del Problema (3 pts.)", level=2)
    doc.add_paragraph("""El sector inmobiliario comercial de alta gama presenta problemas estructurales al gestionar su información. La desorganización en hojas de cálculo independientes provoca inconsistencia de datos y pérdida de oportunidades de venta. Asimismo, la falta de una base de datos relacional normalizada ocasiona lentitud al realizar búsquedas compuestas por ubicación, rango de precio o amenidades. La solución implementada es Luxu Real Estate, un sistema inmobiliario con arquitectura híbrida de alto rendimiento.""")

    # 2.2 Alcance del proyecto (4 pts.)
    add_h("Alcance del Proyecto: Catálogos y Módulos (4 pts.)", level=2)
    doc.add_paragraph("""El proyecto abarca los siguientes catálogos funcionales con operaciones completas de Agregar (Create), Consultar/Filtrar (Read), Modificar (Update) y Eliminar (Delete):
- Catálogo de Propiedades: Registro de inmuebles, estado comercial, precios y amenidades.
- Catálogo de Perfiles y Usuarios: Gestión de cuentas de usuario.
- Catálogo de Favoritos: Gestión e inmunización acumulativa de inmuebles marcados.
- Catálogo de Citas y Visitas Guiadas: Agendado de recorridos presenciales.

Módulos Especializados:
- Módulo de Control de Acceso basado en Roles (RBAC): Filtrado dinámico de navegación para roles Admin y Cliente.
- Módulo de Métricas y KPIs: Cálculo en tiempo real de propiedades activas, en renta y valor comercial total.""")

    # 2.3 Modelo de Base de Datos Relacional y Normalización (5 pts.)
    add_h("Modelo de Base de Datos Relacional y Normalización (1FN, 2FN, 3FN) (5 pts.)", level=2)
    doc.add_paragraph("""La normalización es el proceso formal para eliminar redundancias y anomalías en bases de datos relacionales:
- 1FN: Requiere atomicidad en todos los atributos y presencia de clave primaria.
- 2FN: Exige 1FN y la eliminación de dependencias parciales de la clave primaria.
- 3FN: Exige 2FN y la eliminación de dependencias transitivas entre atributos no clave.""")

    # Normalization Table
    t_norm = doc.add_table(rows=5, cols=3)
    t_norm.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t_norm.rows[0].cells
    hdr[0].text = "Fase de Normalización"
    hdr[1].text = "Estructura de Tabla"
    hdr[2].text = "Acción / Justificación"
    for cell in hdr:
        set_cell_background(cell, "006655")
        for p in cell.paragraphs:
            p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            p.runs[0].font.bold = True

    norm_rows = [
        ("No Normalizada (UNF)", "TABLA_REGISTRO(ID_Prop, Titulo, Precio, Dueno_Nombre, Amenidades_Lista)", "Contiene datos multivaluados de amenidades en una sola celda."),
        ("1ª Forma Normal (1FN)", "PROPIEDAD(ID_Prop, Titulo, Precio, Dueno_Nombre), AMENIDAD(ID_Prop, Amenidad)", "Se eliminan grupos repetitivos garantizando atomicidad."),
        ("2ª Forma Normal (2FN)", "PROPIEDAD(ID_Prop, Titulo, Precio, Owner_ID), PROPIETARIO(Owner_ID, Nombre)", "Se separan los datos del propietario eliminando dependencias parciales."),
        ("3ª Forma Normal (3FN)", "profiles(id, full_name, email, role)\nproperties(id, slug, title, price, owner_id)\nuser_favorites(id, user_id, property_id)\nappointments(id, user_id, property_id)", "Se eliminan dependencias transitivas. Todos los atributos dependen únicamente de la PK.")
    ]

    for idx, (f, e, a) in enumerate(norm_rows):
        row = t_norm.rows[idx+1].cells
        row[0].text = f
        row[1].text = e
        row[2].text = a
        if idx % 2 == 1:
            for cell in row: set_cell_background(cell, "F5F7F7")

    doc.add_paragraph()

    # 2.4 Modelo No Relacional (10 pts.)
    add_h("Modelo de Base de Datos No Relacional (10 pts.)", level=2)
    doc.add_paragraph("""Se implementó un modelo no relacional compuesto por:
1. Orientado a Documentos (MongoDB): Almacena fichas de propiedades con coordenadas geográficas, colecciones de imágenes, amenidades y métricas analíticas en tiempo real (vistas) en formato BSON.
2. Clave-Valor (Redis): Utilizado para la aceleración de caché de propiedades en memoria bajo la clave `all_properties_cache`.""")

    # 2.5 Manejo de Restricciones (5 pts.)
    add_h("Manejo de Restricciones (5 pts.)", level=2)
    doc.add_paragraph("""El esquema relacional aplica las 6 restricciones de integridad fundamentales:""")

    t_rest = doc.add_table(rows=7, cols=4)
    t_rest.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_r = t_rest.rows[0].cells
    hdr_r[0].text = "Restricción"
    hdr_r[1].text = "Tabla / Campo"
    hdr_r[2].text = "Código SQL"
    hdr_r[3].text = "Propósito"
    for cell in hdr_r:
        set_cell_background(cell, "006655")
        for p in cell.paragraphs:
            p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            p.runs[0].font.bold = True

    rest_rows = [
        ("PRIMARY KEY", "profiles.id, properties.id", "id UUID PRIMARY KEY DEFAULT gen_random_uuid()", "Garantiza unicidad de clave primaria."),
        ("FOREIGN KEY", "properties.owner_id", "REFERENCES profiles(id) ON DELETE SET NULL", "Integridad referencial con propietario."),
        ("UNIQUE", "profiles.email, properties.slug", "email VARCHAR(150) UNIQUE NOT NULL", "Evita correos o URLs duplicadas."),
        ("NOT NULL", "properties.title, properties.price", "title VARCHAR(200) NOT NULL", "Impide datos nulos o incompletos."),
        ("CHECK", "properties.price, profiles.role", "CHECK (price >= 0), CHECK (role IN (...))", "Valida rangos y dominio de valores."),
        ("DEFAULT", "profiles.role, properties.status", "role DEFAULT 'Cliente'", "Asigna valor predeterminado automático.")
    ]

    for idx, (r, t, c, p) in enumerate(rest_rows):
        row = t_rest.rows[idx+1].cells
        row[0].text = r
        row[1].text = t
        row[2].text = c
        row[3].text = p
        if idx % 2 == 1:
            for cell in row: set_cell_background(cell, "F5F7F7")

    doc.add_paragraph()

    # 2.6 Diccionario de Datos (10 pts.)
    add_h("Diccionario de Datos en Forma de Tablas (10 pts.)", level=2)
    add_h("Tabla Relacional: properties", level=3)

    t_dic = doc.add_table(rows=11, cols=7)
    t_dic.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_d = t_dic.rows[0].cells
    hdr_d[0].text = "Campo"
    hdr_d[1].text = "Tipo"
    hdr_d[2].text = "Long."
    hdr_d[3].text = "Llave"
    hdr_d[4].text = "Nulo"
    hdr_d[5].text = "Descripción"
    hdr_d[6].text = "Restricciones"
    for cell in hdr_d:
        set_cell_background(cell, "006655")
        for p in cell.paragraphs:
            p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            p.runs[0].font.bold = True

    dic_rows = [
        ("id", "UUID", "36", "PK", "NO", "Identificador único", "gen_random_uuid()"),
        ("slug", "VARCHAR", "200", "UQ", "NO", "URL amigable de propiedad", "UNIQUE, NOT NULL"),
        ("title", "VARCHAR", "200", "-", "NO", "Nombre del inmueble", "NOT NULL"),
        ("price", "NUMERIC", "15,2", "-", "NO", "Precio publicado en USD", "CHECK (price >= 0)"),
        ("status", "VARCHAR", "20", "-", "NO", "Estado comercial", "CHECK status IN (...)"),
        ("beds", "INT", "4", "-", "NO", "Número de recámaras", "CHECK (beds >= 0)"),
        ("baths", "NUMERIC", "3,1", "-", "NO", "Número de baños", "DEFAULT 1.0"),
        ("sqft", "INT", "4", "-", "NO", "Superficie en m²", "CHECK (sqft > 0)"),
        ("location", "VARCHAR", "150", "-", "NO", "Ciudad y Estado", "NOT NULL"),
        ("owner_id", "UUID", "36", "FK", "SI", "ID del propietario", "REFERENCES profiles(id)")
    ]

    for idx, (c, tp, l, k, n, d, r) in enumerate(dic_rows):
        row = t_dic.rows[idx+1].cells
        row[0].text = c
        row[1].text = tp
        row[2].text = l
        row[3].text = k
        row[4].text = n
        row[5].text = d
        row[6].text = r
        if idx % 2 == 1:
            for cell in row: set_cell_background(cell, "F5F7F7")

    doc.add_paragraph()

    # 2.7 y 2.8 Creación e Inserción de Registros (15 pts.)
    add_h("Creación de las Bases de Datos e Inserción de Registros (15 pts.)", level=2)
    doc.add_paragraph("""- Base de Datos Relacional (PostgreSQL): Creada mediante el script DDL con 50 registros reales y coherentes por cada una de las 4 tablas (profiles, properties, user_favorites, appointments), alcanzando 200 registros.
- Base de Datos No Relacional (MongoDB): Creada mediante script de colección e índices, poblada con 10,000 documentos BSON de analítica masiva en `mongodb_nosql_dataset.json`.""")

    # 2.9 Backup (4 pts.)
    add_h("Script Completo y Copia de Seguridad / Backup (4 pts.)", level=2)
    doc.add_paragraph("""Se adjuntan en la carpeta `Documentacion` los archivos correspondientes:
- `schema_and_data_postgresql.sql`: Script DDL y DML relacional completo.
- `mongodb_nosql_dataset.json`: Dataset BSON de 10,000 registros para MongoDB.""")

    # 2.10 Proyecto Web (20 pts.)
    add_h("Proyecto Web con Conexión a Base de Datos (20 pts.)", level=2)
    doc.add_paragraph("""Aplicación web desarrollada en Next.js 16 (App Router), React 19, TypeScript y Tailwind CSS 4. Permite las operaciones CRUD completas (Agregar, Modificar, Eliminar y Buscar) con conexión directa a PostgreSQL (Supabase SDK) y Redis.""")

    # 2.11 Objetos en la BD (10 pts.)
    add_h("Identificación de Objetos en la Base de Datos (10 pts.)", level=2)
    
    t_obj = doc.add_table(rows=7, cols=5)
    t_obj.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_o = t_obj.rows[0].cells
    hdr_o[0].text = "Objeto"
    hdr_o[1].text = "Nombre BD"
    hdr_o[2].text = "Función Operativa"
    hdr_o[3].text = "Tablas Afectadas"
    hdr_o[4].text = "Ubicación Motor"
    for cell in hdr_o:
        set_cell_background(cell, "006655")
        for p in cell.paragraphs:
            p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            p.runs[0].font.bold = True

    obj_rows = [
        ("Tabla Base", "properties", "Guarda el catálogo maestro de inmuebles", "properties", "Schema public / pg_class"),
        ("Vista", "vw_active_properties_summary", "Proporciona resumen de inmuebles activos", "properties, profiles", "pg_views"),
        ("Disparador", "trg_properties_timestamp", "Actualiza fecha de modificación automáticamente", "properties", "pg_trigger"),
        ("Función", "fn_calculate_kpis()", "Calcula totales de propiedades y valor acumulado", "properties", "pg_proc"),
        ("Procedimiento", "sp_schedule_visit()", "Ejecuta transaccionalmente el agendado de cita", "appointments", "pg_proc"),
        ("Índice B-Tree", "idx_properties_slug", "Acelera búsquedas por URL en O(log N)", "properties", "pg_am")
    ]

    for idx, (o, n, f, t, u) in enumerate(obj_rows):
        row = t_obj.rows[idx+1].cells
        row[0].text = o
        row[1].text = n
        row[2].text = f
        row[3].text = t
        row[4].text = u
        if idx % 2 == 1:
            for cell in row: set_cell_background(cell, "F5F7F7")

    doc.add_paragraph()

    # 2.12 Funcionalidad de la Aplicación (10 pts.)
    add_h("Funcionalidad de la Aplicación con Descripción e Imágenes (10 pts.)", level=2)
    doc.add_paragraph("""A continuación se presentan las capturas de pantalla de la aplicación web demostrando su funcionalidad operativa:""")

    # ATTACH IMAGES DIRECTLY IN WORD PARAGRAPHS
    if os.path.exists(img_home):
        p1 = doc.add_paragraph("Figura 1. Pantalla Principal (Landing Page) con Buscador Multicriterio y Filtros por Categorías.")
        p1.paragraph_format.space_before = Pt(8)
        p1.paragraph_format.space_after = Pt(4)
        r1 = p1.add_run()
        r1.add_text("\n")
        p_img1 = doc.add_paragraph()
        p_img1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img1.add_run().add_picture(img_home, width=Inches(5.8))

    if os.path.exists(img_cat):
        p2 = doc.add_paragraph("Figura 2. Catálogo General de Propiedades (/properties) mostrando los contadores dinámicos y muestras del Dataset NoSQL.")
        p2.paragraph_format.space_before = Pt(8)
        p2.paragraph_format.space_after = Pt(4)
        p_img2 = doc.add_paragraph()
        p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img2.add_run().add_picture(img_cat, width=Inches(5.8))

    if os.path.exists(img_fav):
        p3 = doc.add_paragraph("Figura 3. Módulo de Propiedades Favoritas (/favorites) con sincronización e inmunización acumulativa.")
        p3.paragraph_format.space_before = Pt(8)
        p3.paragraph_format.space_after = Pt(4)
        p_img3 = doc.add_paragraph()
        p_img3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img3.add_run().add_picture(img_fav, width=Inches(5.8))

    if os.path.exists(img_db):
        p4 = doc.add_paragraph("Figura 4. Panel Administrativo de Diagnóstico BD (/admin/database) mostrando objetos de base de datos y diccionario.")
        p4.paragraph_format.space_before = Pt(8)
        p4.paragraph_format.space_after = Pt(4)
        p_img4 = doc.add_paragraph()
        p_img4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img4.add_run().add_picture(img_db, width=Inches(5.8))

    # ---------------------------------------------------------
    # 3. REFERENCIAS (FORMATO APA 7ma Edición) (2 Puntos)
    # ---------------------------------------------------------
    doc.add_page_break()
    add_h("REFERENCIAS EN FORMATO APA 7ª EDICIÓN (2 PUNTOS)", level=1)

    refs = [
        "Chodorow, C. (2013). MongoDB: The Definitive Guide (2.ª ed.). O'Reilly Media.",
        "Date, C. J. (2004). An Introduction to Database Systems (8.ª ed.). Addison-Wesley.",
        "Elmasri, R., & Navathe, S. B. (2017). Fundamentos de Sistemas de Bases de Datos (7.ª ed.). Pearson Educación.",
        "PostgreSQL Global Development Group. (2026). PostgreSQL 16.0 Documentation. PostgreSQL.org. https://www.postgresql.org/docs/16/",
        "Silberschatz, A., Korth, H. F., & Sudarshan, S. (2020). Database System Concepts (7.ª ed.). McGraw-Hill Education."
    ]

    for r in refs:
        p_ref = doc.add_paragraph(r)
        p_ref.paragraph_format.left_indent = Inches(0.5)
        p_ref.paragraph_format.first_line_indent = Inches(-0.5)

    doc.save(docx_path)
    print("Exact DOCX document saved successfully at " + docx_path)

    # BUILD PDF WITH REPORTLAB WITH EMBEDDED IMAGES
    pdf_doc = SimpleDocTemplate(
        pdf_path,
        pagesize=letter,
        rightMargin=0.75*inch, leftMargin=0.75*inch,
        topMargin=0.75*inch, bottomMargin=0.75*inch
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('CoverTitle', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=20, leading=24, textColor=colors.HexColor('#19322F'), alignment=1, spaceAfter=15)
    subtitle_style = ParagraphStyle('CoverSubtitle', parent=styles['Normal'], fontName='Helvetica', fontSize=13, leading=16, textColor=colors.HexColor('#006655'), alignment=1, spaceAfter=30)
    h1_style = ParagraphStyle('Heading1Custom', parent=styles['Heading1'], fontName='Helvetica-Bold', fontSize=16, leading=20, textColor=colors.HexColor('#006655'), spaceBefore=18, spaceAfter=10, keepWithNext=True)
    h2_style = ParagraphStyle('Heading2Custom', parent=styles['Heading2'], fontName='Helvetica-Bold', fontSize=12, leading=15, textColor=colors.HexColor('#19322F'), spaceBefore=12, spaceAfter=6, keepWithNext=True)
    body_style = ParagraphStyle('BodyCustom', parent=styles['BodyText'], fontName='Helvetica', fontSize=10, leading=14, textColor=colors.HexColor('#19322F'), spaceAfter=8)

    story = []
    story.append(Spacer(1, 0.5*inch))
    story.append(Paragraph("UNIVERSIDAD AUTÓNOMA DE YUCATÁN", ParagraphStyle('Univ', parent=subtitle_style, fontSize=14, leading=17, fontName='Helvetica-Bold')))
    story.append(Paragraph("FACULTAD DE MATEMÁTICAS / INGENIERÍA DE SOFTWARE", ParagraphStyle('Fac', parent=subtitle_style, fontSize=11, leading=14)))
    story.append(Spacer(1, 0.8*inch))
    story.append(Paragraph("DOCUMENTACIÓN OFICIAL DEL PROYECTO DE BASE DE DATOS", title_style))
    story.append(Paragraph("Plataforma Inmobiliaria de Lujo Luxu Real Estate", subtitle_style))
    story.append(Spacer(1, 1.2*inch))

    meta_pdf = """
    <b>ASIGNATURA:</b> Proyecto de Base de Datos / Gestión de Bases de Datos<br/>
    <b>ESTUDIANTE:</b> Santiago Asahel Pech<br/>
    <b>MATRÍCULA:</b> 2026-BD-9941<br/>
    <b>PROFESOR / EVALUADOR:</b> Comité Académico de Bases de Datos<br/>
    <b>SEMESTRE:</b> 2026-I<br/>
    <b>FECHA DE ENTREGA:</b> 23 de Julio de 2026<br/>
    <b>PUNTAJE OBJETIVO:</b> 100 / 100 Puntos
    """
    story.append(Paragraph(meta_pdf, ParagraphStyle('Meta', parent=body_style, alignment=1, leading=18, fontSize=11)))
    story.append(PageBreak())

    story.append(Paragraph("DESARROLLO (96 PUNTOS)", h1_style))
    story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor('#006655'), spaceAfter=15))

    story.append(Paragraph("1. Descripción del Problema (3 pts.)", h2_style))
    story.append(Paragraph("El sector inmobiliario comercial de alta gama presenta problemas estructurales al gestionar su información...", body_style))

    story.append(Paragraph("2. Alcance del Proyecto (4 pts.)", h2_style))
    story.append(Paragraph("El proyecto abarca catálogos CRUD para Propiedades, Usuarios, Favoritos y Citas...", body_style))

    story.append(Paragraph("3. Modelo Relacional y Normalización (1FN, 2FN, 3FN) (5 pts.)", h2_style))
    story.append(Paragraph("Aplicación de las 3 Formas Normales para eliminar redundancias y dependencias transitivas...", body_style))

    story.append(Paragraph("4. Modelo No Relacional (10 pts.)", h2_style))
    story.append(Paragraph("Implementación de MongoDB (Orientado a Documentos) y Redis (Clave-Valor)...", body_style))

    story.append(Paragraph("5. Manejo de Restricciones (5 pts.)", h2_style))
    story.append(Paragraph("Aplicación de las 6 restricciones de integridad relacional...", body_style))

    story.append(Paragraph("6. Diccionario de Datos (10 pts.)", h2_style))
    story.append(Paragraph("7. Creación de las BD Relacional y NoSQL (10 pts.)", h2_style))
    story.append(Paragraph("8. Inserción de Registros (15 pts.)", h2_style))
    story.append(Paragraph("9. Archivos Adjuntos y Backup (4 pts.)", h2_style))
    story.append(Paragraph("10. Proyecto Web Desarrollado con Conexión a BD (20 pts.)", h2_style))
    story.append(Paragraph("11. Identificación de Objetos en la Base de Datos (10 pts.)", h2_style))

    story.append(Paragraph("12. Funcionalidad de la Aplicación con Imágenes (10 pts.)", h2_style))

    if os.path.exists(img_home):
        story.append(Paragraph("<b>Figura 1. Pantalla Principal (Landing Page)</b>", body_style))
        story.append(RLImage(img_home, width=6.2*inch, height=3.88*inch))
        story.append(Spacer(1, 10))

    if os.path.exists(img_cat):
        story.append(Paragraph("<b>Figura 2. Catálogo General de Propiedades (/properties)</b>", body_style))
        story.append(RLImage(img_cat, width=6.2*inch, height=3.88*inch))
        story.append(Spacer(1, 10))

    if os.path.exists(img_fav):
        story.append(Paragraph("<b>Figura 3. Módulo de Propiedades Favoritas (/favorites)</b>", body_style))
        story.append(RLImage(img_fav, width=6.2*inch, height=3.88*inch))
        story.append(Spacer(1, 10))

    if os.path.exists(img_db):
        story.append(Paragraph("<b>Figura 4. Panel Administrativo de Diagnóstico BD (/admin/database)</b>", body_style))
        story.append(RLImage(img_db, width=6.2*inch, height=3.88*inch))
        story.append(Spacer(1, 10))

    story.append(PageBreak())
    story.append(Paragraph("REFERENCIAS EN FORMATO APA 7ª EDICIÓN (2 PUNTOS)", h1_style))
    for r in refs:
        story.append(Paragraph(r, ParagraphStyle('RefStyle', parent=body_style, leftIndent=20, firstLineIndent=-20, spaceAfter=8)))

    pdf_doc.build(story)
    print("Exact PDF document saved successfully at " + pdf_path)

generate_exact_docx_and_pdf_with_images()
