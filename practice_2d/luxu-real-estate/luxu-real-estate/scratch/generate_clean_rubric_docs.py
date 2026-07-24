import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image as RLImage
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch

def build_rubric_clean_documentation():
    docx_path = r"c:\home\practice_2d\luxu-real-estate\luxu-real-estate\Documentacion\Documentacion_Proyecto_Base_de_Datos.docx"
    pdf_path = r"c:\home\practice_2d\luxu-real-estate\luxu-real-estate\Documentacion\Documentacion_Proyecto_Base_de_Datos.pdf"
    scratch = r"c:\home\practice_2d\luxu-real-estate\luxu-real-estate\scratch"

    doc = docx.Document()
    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)

    # Helper styling functions
    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(0, 102, 85)

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(25, 50, 47)

    def add_p(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(text)
        r.font.size = Pt(10.5)

    # 1. PORTADA
    p_cover = doc.add_paragraph()
    p_cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_univ = p_cover.add_run("UNIVERSIDAD AUTÓNOMA DE YUCATÁN\nFACULTAD DE MATEMÁTICAS / INGENIERÍA DE SOFTWARE\n\n")
    r_univ.bold = True
    r_univ.font.size = Pt(13)
    r_univ.font.color.rgb = RGBColor(0, 102, 85)

    r_title = p_cover.add_run("DOCUMENTACIÓN TÉCNICA Y EVIDENCIAS DEL PROYECTO DE BASE DE DATOS\n")
    r_title.bold = True
    r_title.font.size = Pt(15)
    r_title.font.color.rgb = RGBColor(25, 50, 47)

    r_sub = p_cover.add_run("SISTEMA DE GESTIÓN INMOBILIARIA HÍBRIDO: LUXU REAL ESTATE\n\n")
    r_sub.bold = True
    r_sub.font.size = Pt(12)
    r_sub.font.color.rgb = RGBColor(0, 102, 85)

    r_meta = p_cover.add_run(
        "Asignatura: Proyecto de Base de Datos\n"
        "Estudiante: Santiago Asahel Pech\n"
        "Evaluador: Comité Académico de Bases de Datos\n"
        "Tecnologías: Next.js 16, React 19, PostgreSQL (Supabase 3FN), MongoDB Atlas (10,000 Registros), Redis\n"
        "Fecha: Julio 2026\n\n"
    )
    r_meta.font.size = Pt(10.5)

    # 2. DESARROLLO
    add_h1("DESARROLLO DE LA BASE DE DATOS Y SISTEMA")

    add_h2("1. Descripción del Problema")
    add_p(
        "El sector inmobiliario de alta gama requiere un manejo de alto rendimiento. Tradicionalmente sufre de desorganización, "
        "consultas lentas bajo filtros estrictos y falta de control de acceso. La plataforma Luxu Real Estate soluciona esto "
        "implementando una arquitectura híbrida: PostgreSQL (Supabase) para transacciones relacionales normalizadas en 3FN, "
        "MongoDB Atlas para analítica masiva NoSQL con 10,000 documentos BSON, y Redis para aceleración en caché."
    )

    add_h2("2. Alcance del Proyecto (Catálogos CRUD y Módulos)")
    add_p(
        "Catálogos desarrollados con operaciones básicas de Agregar, Modificar, Eliminar y Buscar:\n"
        "• Catálogo de Propiedades (Venta y Renta): Gestión integral de listados con imágenes y filtros multicriterio.\n"
        "• Catálogo de Usuarios: Directorio con asignación de Roles (Admin, Agente, Cliente).\n"
        "• Catálogo de Favoritos: Guardado dinámico acumulativo por usuario.\n"
        "• Módulo de Citas y Visitas Guiadas: Agendado de recorridos presenciales.\n"
        "• Módulo de Control de Acceso por Roles (RBAC) y Panel Diagnóstico de Base de Datos."
    )

    add_h2("3. Modelo de Base de Datos Relacional y Normalización (1FN, 2FN, 3FN)")
    add_p(
        "Proceso de normalización aplicado:\n"
        "• 1FN: Eliminación de atributos multivaluados y asignación de identificadores primarios (UUID).\n"
        "• 2FN: Eliminación de dependencias parciales separando datos de usuarios y publicaciones.\n"
        "• 3FN: Eliminación de dependencias transitivas. Tablas resultantes con 50 registros reales por tabla: "
        "profiles, properties, user_favorites, appointments."
    )

    add_h2("4. Modelo No Relacional MongoDB Atlas (10,000 Registros)")
    add_p(
        "Colección 'properties_nosql' estructurada en MongoDB Atlas con 10,000 documentos BSON. "
        "Soporta esquemas flexibles, arreglos dinámicos de amenidades e índices geoespaciales 2DSphere y Full-Text Search."
    )

    add_h2("5. Manejo de Restricciones de Integridad")
    add_p(
        "Implementación de las 6 restricciones principales:\n"
        "• PRIMARY KEY: gen_random_uuid() en profiles.id y properties.id.\n"
        "• FOREIGN KEY: owner_id REFERENCES profiles(id) ON DELETE SET NULL.\n"
        "• UNIQUE: Restricción única en email y slug.\n"
        "• NOT NULL: Obligatoriedad en title, price, y status.\n"
        "• CHECK: price >= 0, beds >= 0, role IN ('Admin','Agente','Cliente').\n"
        "• DEFAULT: Valores por defecto en role ('Cliente') y timestamps."
    )

    add_h2("6. Diccionario de Datos y Objetos de la Base de Datos")
    add_p(
        "Estructura relacional y NoSQL documentada. Se identifican objetos clave en PostgreSQL y MongoDB:\n"
        "• Tablas Base: properties, profiles, user_favorites, appointments.\n"
        "• Vistas (Views): vw_active_properties (filtro de inmuebles activos).\n"
        "• Disparadores (Triggers): trg_update_timestamp (actualización de modificación).\n"
        "• Procedimientos y Funciones: fn_calculate_kpis(), sp_schedule_visit().\n"
        "• Índices: idx_prop_slug (B-Tree) e idx_spatial_coords (2DSphere NoSQL)."
    )

    add_h2("7. Creación de las Bases de Datos y Scripts (.sql y Backup)")
    add_p(
        "Se adjunta en la carpeta de entregables el script completo DDL + DML con 50 registros por tabla (schema_and_data_postgresql.sql) "
        "y la copia de seguridad/dataset masivo con 10,000 registros (mongodb_nosql_dataset.json)."
    )

    # 3. EVIDENCIAS FOTOGRÁFICAS ORGANIZADAS
    add_h1("EVIDENCIAS DE FUNCIONALIDAD DEL SISTEMA Y BASES DE DATOS")

    evidencias = [
        ('ev_panel_control.png', 'Figura 1. Módulo Panel de Control (/admin/properties) - Operaciones CRUD'),
        ('ev_vender.png', 'Figura 2. Formulario Agregar Inmueble (/admin/properties/add) - Operación Create'),
        ('ev_usuarios.png', 'Figura 3. Directorio de Usuarios y Roles (/admin/users) - Gestión RBAC'),
        ('ev_comprar.png', 'Figura 4. Catálogo Comprar (/properties?type=buy) - Búsqueda y Filtros'),
        ('ev_rentar.png', 'Figura 5. Catálogo Rentar (/properties?type=rent) - Inmuebles en Alquiler'),
        ('ev_favoritas.png', 'Figura 6. Módulo Propiedades Favoritas (/favorites) - Persistencia de Usuario'),
        ('ev_perfil.png', 'Figura 7. Perfil de Usuario (/profile) - Métricas e Historial'),
        ('ev_agendar_visita.png', 'Figura 8. Módulo Agendar Visita Presencial (/schedule-visit)'),
        ('cap_supabase_tables.png', 'Figura 9. Tablas Relacionales PostgreSQL en Supabase (50 Tuplas/Tabla en 3FN)'),
        ('ev_base_datos.png', 'Figura 10. Base de Datos MongoDB Atlas (10,000 Documentos NoSQL Almacenados)')
    ]

    for img_name, label in evidencias:
        img_full = os.path.join(scratch, img_name)
        if os.path.exists(img_full):
            p_lbl = doc.add_paragraph()
            p_lbl.paragraph_format.space_before = Pt(10)
            p_lbl.paragraph_format.space_after = Pt(3)
            r_lbl = p_lbl.add_run(label)
            r_lbl.bold = True
            r_lbl.font.size = Pt(11)
            r_lbl.font.color.rgb = RGBColor(0, 102, 85)

            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_after = Pt(10)
            p_img.add_run().add_picture(img_full, width=Inches(6.2))

    # 4. REFERENCIAS APA 7ma EDICIÓN
    add_h1("REFERENCIAS (FORMATO APA 7ma Edición)")
    add_p("1. Elmasri, R., & Navathe, S. B. (2017). Fundamentos de Sistemas de Bases de Datos (7.ª ed.). Pearson Educación.")
    add_p("2. Date, C. J. (2004). An Introduction to Database Systems (8.ª ed.). Addison-Wesley.")
    add_p("3. Silberschatz, A., Korth, H. F., & Sudarshan, S. (2020). Database System Concepts (7.ª ed.). McGraw-Hill Education.")
    add_p("4. Chodorow, C. (2013). MongoDB: The Definitive Guide (2.ª ed.). O'Reilly Media.")
    add_p("5. PostgreSQL Global Development Group. (2026). PostgreSQL 16.0 Documentation. https://www.postgresql.org/docs/16/")

    doc.save(docx_path)
    print("Clean DOCX generated successfully matching strictly rubric items!")

    # PDF BUILD
    doc_pdf = SimpleDocTemplate(pdf_path, pagesize=letter, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()
    normal = styles['Normal']

    title_style = ParagraphStyle('PDFTitle', parent=normal, fontName='Helvetica-Bold', fontSize=13, leading=16, textColor=colors.HexColor('#006655'), alignment=1, spaceAfter=10)
    h2_style = ParagraphStyle('PDFH2', parent=normal, fontName='Helvetica-Bold', fontSize=11, leading=14, textColor=colors.HexColor('#19322F'), spaceBefore=8, spaceAfter=4)
    body_style = ParagraphStyle('PDFBody', parent=normal, fontName='Helvetica', fontSize=9.5, leading=13, textColor=colors.HexColor('#333333'), spaceAfter=6)
    lbl_style = ParagraphStyle('PDFLbl', parent=normal, fontName='Helvetica-Bold', fontSize=10, leading=13, textColor=colors.HexColor('#006655'), spaceBefore=8, spaceAfter=4)

    story = [
        Paragraph("UNIVERSIDAD AUTÓNOMA DE YUCATÁN - INGENIERÍA DE SOFTWARE", title_style),
        Paragraph("DOCUMENTACIÓN OFICIAL Y EVIDENCIAS DE PROYECTO DE BASE DE DATOS", title_style),
        Spacer(1, 8),
        Paragraph("Desarrollo de la Base de Datos Relacional (3FN) y NoRelacional (10,000 Registros)", h2_style),
        Paragraph("El proyecto Luxu Real Estate implementa una arquitectura híbrida con PostgreSQL (Supabase) en 3FN y MongoDB Atlas con 10,000 documentos NoSQL para analítica masiva.", body_style),
        Spacer(1, 8),
        Paragraph("Evidencias del Sistema y Bases de Datos (Supabase & MongoDB)", h2_style)
    ]

    for img_name, label in evidencias:
        img_full = os.path.join(scratch, img_name)
        if os.path.exists(img_full):
            story.append(Paragraph(label, lbl_style))
            story.append(RLImage(img_full, width=6.5*inch, height=3.8*inch))
            story.append(Spacer(1, 8))

    doc_pdf.build(story)
    print("Clean PDF generated successfully matching strictly rubric items!")

if __name__ == '__main__':
    build_rubric_clean_documentation()
